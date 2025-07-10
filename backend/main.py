import os
import uuid
import zipfile
import shutil
import re
from pathlib import Path

from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, HTMLResponse, Response

from docx import Document
from docx.shared import Inches, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import load_workbook

# --- FastAPI 應用程式設定 ---
app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"], allow_credentials=True,
    allow_methods=["*"], allow_headers=["*"],
)

# --- 路徑與常數設定 ---
BASE_DIR = Path(__file__).resolve().parent.parent
FRONTEND_DIR = BASE_DIR / "frontend"
TEMP_DIR = BASE_DIR / "temp"
TEMP_DIR.mkdir(exist_ok=True)
MAIN_CHAPTER_NUMBER = "1"

# --- 靜態檔案與首頁路由 ---
app.mount("/static", StaticFiles(directory=FRONTEND_DIR), name="static")

@app.get("/", response_class=HTMLResponse)
async def root():
    index_path = FRONTEND_DIR / "index.html"
    if not index_path.exists():
        raise HTTPException(status_code=404, detail="index.html not found")
    return HTMLResponse(index_path.read_text(encoding="utf-8"))

@app.get("/favicon.ico", include_in_schema=False)
async def favicon(): return Response(status_code=204)

# --- 核心輔助函式 ---
def set_font_for_run(run):
    run.font.name = 'Times New Roman'
    run.font.element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')

def add_paragraph_with_font(doc, text, style=None):
    p = doc.add_paragraph(text, style=style)
    for r in p.runs: set_font_for_run(r)
    return p

def add_heading_with_font(doc, text, level):
    h = doc.add_heading(text, level=level)
    for r in h.runs: set_font_for_run(r)
    return h

# --- 讀 Excel ---
def parse_excel_config(xlsx_path: Path):
    wb = load_workbook(xlsx_path, read_only=True)
    ws = wb.worksheets[0]
    sections = []
    for idx, (col_a, col_b) in enumerate(ws.iter_rows(min_row=2, max_col=2, values_only=True), 1):
        if not col_a:
            continue
        title_raw = str(col_a).strip()
        query_var = str(col_b).strip() if col_b else "（未提供）"

        delimiter = "功能代碼:"
        folder_id = ""
        if delimiter in title_raw:
            parts = title_raw.split(delimiter, 1)
            title = parts[0].strip().removesuffix(',').removesuffix('，').strip()
            folder_id = parts[1].strip()
        else:
            title = title_raw

        sections.append({
            'folder': folder_id,
            'title':  title,
            'query_condition': query_var,
            'base_number': f"{MAIN_CHAPTER_NUMBER}.{idx}"
        })
    return sections

# ---  append_section   ---
def append_section(doc: Document, section_info: dict, image_folder: str):
    base_number       = section_info.get("base_number", "?")
    section_title     = section_info.get("title", "未命名標題")
    section_folder_name = section_info.get("folder", "未知資料夾")
    query_condition   = section_info.get("query_condition", "（未提供）")

    try:
        images = sorted(
            os.path.join(image_folder, f) for f in os.listdir(image_folder)
            if f.lower().endswith(('.png', '.jpg', '.jpeg'))
        )
    except FileNotFoundError:
        add_paragraph_with_font(doc, f"警告：找不到與 {section_folder_name} 對應的圖片資料夾。")
        images = []

    uniform_width_cm = 17.98

    def add_picture_with_uniform_width(idx):
        if idx < len(images):
            try:
                # 使用 Cm 物件來設定所有圖片的寬度
                width = Cm(uniform_width_cm)
                doc.add_picture(images[idx], width=width)
                # 將包含圖片的段落置中
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                p = add_paragraph_with_font(doc, f"（警告：處理第 {idx+1} 張圖片時發生錯誤: {e}）")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p = add_paragraph_with_font(doc, f"（警告：此處缺少第 {idx+1} 張圖片）")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- 文件內容生成 ---

    # 主標題
    doc.add_page_break()
    add_heading_with_font(doc, f"{section_title}，功能代碼:{section_folder_name}", level=2)

    # 查詢功能 > 查詢所有資料
    add_paragraph_with_font(doc,
        f"在左側的功能選單中，點選「開關帳及TMP資料查詢放行」，接著點選「IFRS15查詢及放行」，最後再點選「{section_title}」，即可開啟視窗。")
    add_picture_with_uniform_width(0)  # 圖片1

    add_heading_with_font(doc, "查詢功能", level=3)

    add_heading_with_font(doc, "查詢所有資料", level=4)
    add_paragraph_with_font(doc, "不輸入查詢條件，直接點擊查詢按鈕即可查詢所有資料。")
    add_picture_with_uniform_width(1)  # 圖片2

    # 查詢特定資料
    add_heading_with_font(doc, "查詢特定資料", level=4)
    add_paragraph_with_font(doc, f"查詢條件為：{query_condition}")
    add_paragraph_with_font(doc, "先輸入查詢條件，再點擊查詢按鈕即可查詢特定資料。")
    add_picture_with_uniform_width(2)  # 圖片3
    doc.add_page_break()        # ★★★ 分頁位置保持不變

    # 查詢結果
    add_heading_with_font(doc, "查詢結果", level=4)
    add_picture_with_uniform_width(3)  # 圖片4

    # 核可功能
    add_heading_with_font(doc, "核可功能", level=3)
    add_paragraph_with_font(doc, "若要執行核可功能，需先點擊查詢，確認有資料且資料無誤後再點擊核可按鈕。")

    add_heading_with_font(doc, "已查詢且有資料，點擊核可按鈕", level=4)
    add_picture_with_uniform_width(4)  # 圖片5
    doc.add_page_break()        # ★★★ 分頁位置保持不變

    add_paragraph_with_font(doc, "跳出確認訊息。")
    add_picture_with_uniform_width(5)  # 圖片6
    add_paragraph_with_font(doc, "點擊確認，完成核可後顯示核可者、核可時間")
    add_picture_with_uniform_width(6)  # 圖片7
    doc.add_page_break()        # ★★★ 分頁位置保持不變

    add_heading_with_font(doc, "未查詢或無資料，點擊核可按鈕", level=4)
    add_paragraph_with_font(doc, "若未查詢或無資料就直接點擊核可按鈕，將會跳出警告。")
    add_picture_with_uniform_width(7)  # 圖片8

    add_heading_with_font(doc, "退回", level=3)
    add_paragraph_with_font(doc, "若已完成核可要執行退回功能，可直接點擊退回按鈕。")
    add_picture_with_uniform_width(8)  # 圖片9
    doc.add_page_break()        # ★★★ 分頁位置保持不變
    
    add_paragraph_with_font(doc, "跳出確認訊息。")
    add_picture_with_uniform_width(9)  # 圖片10
    add_paragraph_with_font(doc, "點擊確認，退回後不再顯示核可者、核可時間")
    add_picture_with_uniform_width(10) # 圖片11

# --- API 端點 ---
@app.post("/upload-and-process/")
async def upload_and_process(
    main_docx:   UploadFile = File(...),
    config_file: UploadFile = File(...),  # Excel
    images_zip:  UploadFile = File(...)
):
    sid_path = TEMP_DIR / str(uuid.uuid4()); sid_path.mkdir()

    # 存檔
    main_path = sid_path / main_docx.filename
    cfg_path  = sid_path / config_file.filename
    zip_path  = sid_path / "images.zip"
    for up, p in ((main_docx, main_path), (config_file, cfg_path), (images_zip, zip_path)):
        with p.open("wb") as f: shutil.copyfileobj(up.file, f)

    # 解析 Excel
    sections = parse_excel_config(cfg_path)
    if not sections:
        raise HTTPException(400, "Excel 沒有有效資料列")

    # 解壓圖
    img_root = sid_path / "unzipped_images"
    with zipfile.ZipFile(zip_path) as zf: zf.extractall(img_root)
    subs = list(img_root.iterdir())
    if len(subs) == 1 and subs[0].is_dir(): img_root = subs[0]

    # 產生文件
    doc = Document(main_path)
    for s in sections:
        append_section(doc, s, str(img_root / s['folder']))
    out_name = f"processed_{main_docx.filename}"
    out_path = sid_path / out_name
    doc.save(out_path)

    return {"message": "處理完成！", "output_filename": out_name,
            "download_filename": f"{sid_path.name}/{out_name}"}

# --- 下載 ---
@app.get("/download/{sid}/{fname}")
async def download_file(sid: str, fname: str):
    p = TEMP_DIR / sid / fname
    if not p.exists(): raise HTTPException(404, "找不到檔案")
    return FileResponse(str(p), filename=fname,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")